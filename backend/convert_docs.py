import sys
sys.path.insert(0, r"C:\Users\Acer\AppData\Roaming\Python\Python313\site-packages")

from docx import Document
import os

docs_dir = r"d:\asset-sharing-system\backend\Backend Asset Sharing"
output_dir = r"d:\asset-sharing-system\backend"

files = [
    "BRD - Business Requirements Document",
    "SyRS - System requirements specification",
    "ARCHITECTURE DESIGN DOCUMENT",
    "DATABASE DESIGN DOCUMENT",
]

for fname in files:
    path = os.path.join(docs_dir, fname + ".docx")
    out_path = os.path.join(output_dir, fname + ".txt")
    try:
        doc = Document(path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    content.append(row_text)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(content))
        print(f"Converted: {fname} ({len(content)} lines)")
    except Exception as e:
        print(f"Error converting {fname}: {e}")

print("Done")
